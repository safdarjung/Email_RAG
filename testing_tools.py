#!/usr/bin/env python3
from __future__ import print_function
import os
import re
import json
import pickle
import base64
import csv
import time
import logging
import tempfile
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText
from typing import Any, Dict, List, Optional, Tuple, ClassVar
import contextlib
import io

from dotenv import load_dotenv
from pydantic import Field
from crewai import LLM, Agent
from crewai.tools.base_tool import BaseTool

from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request

from PyPDF2 import PdfReader
import openpyxl
import docx

from litellm.exceptions import RateLimitError

try:
    from PIL import Image, ImageTk
    import pytesseract
except ImportError:
    pytesseract = None

try:
    from pdf2image import convert_from_path
except ImportError:
    convert_from_path = None

import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

# For fuzzy matching of file names
try:
    from rapidfuzz import fuzz
except ImportError:
    fuzz = None

# For in-memory embedding-based retrieval
try:
    from sentence_transformers import SentenceTransformer, util
    EMBEDDING_MODEL = SentenceTransformer('all-MiniLM-L6-v2')
except ImportError:
    EMBEDDING_MODEL = None

# Load environment variables and set scopes
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

###############################################################################
# Global LLM instance (unified_llm)
###############################################################################
unified_llm = LLM(
    model="gemini/gemini-2.0-flash",  # Update with your actual model if needed.
    api_key=GEMINI_API_KEY,
    temperature=0.1
)

###############################################################################
# Function: llm_predict_with_backoff
###############################################################################
def llm_predict_with_backoff(prompt: str, max_retries: int = 5, initial_delay: float = 1.0) -> str:
    """Calls the unified LLM with exponential backoff."""
    delay = initial_delay
    for attempt in range(max_retries):
        try:
            time.sleep(2.0)
            response = unified_llm.call(prompt)
            return response
        except RateLimitError as e:
            logging.warning("Rate limit error on attempt %d/%d: %s", attempt + 1, max_retries, e)
            time.sleep(delay)
            delay *= 2
        except Exception as e:
            logging.error("Error during LLM call: %s", e)
            return ""
    raise RateLimitError("Exceeded maximum retry attempts for LLM call")

###############################################################################
# Function: extract_text_from_file
###############################################################################
def extract_text_from_file(file_path: str) -> str:
    """
    Extracts text from a file (PDF, Excel, Word, CSV, image, or plain text).
    """
    text = ""
    fname = file_path.lower()
    try:
        if fname.endswith('.pdf'):
            with open(file_path, 'rb') as file_content:
                reader = PdfReader(file_content)
                extracted = ""
                for page in reader.pages:
                    pg_text = page.extract_text()
                    if pg_text:
                        extracted += pg_text + "\n"
            if len(extracted.strip()) < 50 and convert_from_path:
                pages = convert_from_path(file_path, dpi=200)
                if pytesseract:
                    for p in pages:
                        extracted += pytesseract.image_to_string(p) + "\n"
            text = extracted
        elif fname.endswith(('.xlsx', '.xls')):
            workbook = openpyxl.load_workbook(file_path)
            sheet = workbook.active
            rows = []
            for row in sheet.iter_rows(values_only=True):
                rows.append("\t".join(str(cell) if cell is not None else "" for cell in row))
            text = "\n".join(rows)
        elif fname.endswith('.docx'):
            document = docx.Document(file_path)
            text = "\n".join(p.text for p in document.paragraphs)
        elif fname.endswith('.csv'):
            with open(file_path, 'r', encoding='utf-8') as file_content:
                reader = csv.reader(file_content)
                rows = []
                for row in reader:
                    rows.append("\t".join(row))
                text = "\n".join(rows)
        elif fname.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
            if pytesseract:
                image = Image.open(file_path)
                text = pytesseract.image_to_string(image)
            else:
                text = "[No OCR support installed]"
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        return text
    except Exception as e:
        logging.error("Error extracting text from %s: %s", file_path, e)
        return f"Error extracting text from {file_path}: {e}"

###############################################################################
# Function: interpret_user_query
###############################################################################
def interpret_user_query(query: str) -> Dict[str, Any]:
    """
    Uses the LLM to convert a natural language query into a Gmail search query.
    Returns a dict with keys: gmail_query, fetch_mode, and has_attachments.
    """
    prompt = (
        f"Interpret the following user query and convert it into a valid Gmail search query by understanding the user's intent. "
        f"Use one of the following formats as appropriate:\n"
        f"  - from:example@gmail.com has:attachment\n"
        f"  - to:recipient@gmail.com after:2023/01/01 before:2023/12/31\n"
        f"  - subject:(\"meeting notes\" OR \"project update\")\n"
        f"  - is:unread label:work\n"
        f"  - \"important information\" -subject:\"important information\"\n"
        f"  - is:starred subject:report\n"
        f"  - in:spam from:spammer@example.com\n"
        f"  - has:attachment filename:pdf\n"
        f"  - \"follow up\" before:2023/01/01\n\n"
        f"Also, decide if the fetch_mode should be 'summary', 'full', or 'analyze', and set has_attachments to true if attachments are expected.\n\n"
        f"User Query: \"{query}\"\n"
        f"Return only the JSON with keys 'gmail_query', 'fetch_mode', and 'has_attachments'."
    )
    result = llm_predict_with_backoff(prompt)
    try:
        result = result.replace("```json", "").replace("```", "").strip()
        data = json.loads(result)
        if not any(op in data.get("gmail_query", "").lower() for op in ["from:", "to:", "after:", "before:", "subject:", "is:", "label:", "has:"]) and "\"" not in data.get("gmail_query", ""):
            if data.get("gmail_query", "").strip():
                data["gmail_query"] = f"\"{data['gmail_query']}\""
            else:
                raise ValueError("Interpreted query does not appear valid.")
        return data
    except Exception as e:
        logging.error("Error interpreting query: %s", e)
        return {
            "gmail_query": f"\"{query}\"",
            "fetch_mode": "summary",
            "has_attachments": "attachment" in query.lower()
        }

###############################################################################
# Function: fetch_gmail_data
###############################################################################
def fetch_gmail_data(query: Optional[str] = None, batch_size: int = 50) -> List[Dict[str, Any]]:
    """
    Uses Gmail API (with OAuth) to fetch email messages matching the query.
    Returns a list of message dictionaries.
    """
    creds = None
    token_path = 'token.pickle'
    if os.path.exists(token_path):
        with open(token_path, 'rb') as token:
            creds = pickle.load(token)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        with open(token_path, 'wb') as token:
            pickle.dump(creds, token)
    service = build('gmail', 'v1', credentials=creds)
    fetched_messages = []
    page_token = None
    while len(fetched_messages) < batch_size:
        if query:
            results = service.users().messages().list(
                userId='me',
                q=query,
                maxResults=min(batch_size, 100),
                pageToken=page_token
            ).execute()
        else:
            results = service.users().messages().list(
                userId='me',
                labelIds=['INBOX'],
                maxResults=min(batch_size, 100),
                pageToken=page_token
            ).execute()
        messages = results.get('messages', [])
        page_token = results.get('nextPageToken')
        for message in messages:
            msg = service.users().messages().get(
                userId='me',
                id=message['id'],
                format='full'
            ).execute()
            fetched_messages.append(msg)
            if len(fetched_messages) >= batch_size:
                break
        if not page_token:
            break
        time.sleep(1)
    return fetched_messages

###############################################################################
# Function: edit_file_content
###############################################################################
def edit_file_content(initial_content: str, history: str = "") -> str:
    """
    Opens a dialog for natural language editing or data analysis.
    The user's instruction is passed to the unified LLM.
    The returned text may include edits, analysis, or instructions for
    visualization. If code is returned, it can be enclosed in triple backticks.
    """
    edit_win = tk.Toplevel()
    edit_win.title("Edit File Content or Analyze Data")
    edit_win.geometry("600x400")

    tk.Label(edit_win, text="Current Content (read-only):").pack(anchor="w", padx=10, pady=(10, 0))
    content_text = ScrolledText(edit_win, wrap=tk.WORD, height=10)
    content_text.pack(fill="both", expand=True, padx=10, pady=5)
    content_text.insert("1.0", initial_content)
    content_text.configure(state=tk.DISABLED)

    tk.Label(edit_win, text="Enter your editing or analysis instruction:").pack(anchor="w", padx=10, pady=(10, 0))
    instruction_entry = tk.Text(edit_win, wrap=tk.WORD, height=4)
    instruction_entry.pack(fill="x", padx=10, pady=5)

    result = {"edited": initial_content}

    def submit_edit():
        instruction = instruction_entry.get("1.0", tk.END).strip()
        if not instruction:
            messagebox.showerror("Error", "Instruction cannot be empty.")
            return

        # Prompt the unified LLM to either edit text or perform data analysis/visualization.
        prompt = (
            "You are an advanced text editor and data analyst. "
            "The user will provide a piece of text or data content plus an instruction. "
            "If the user wants to edit the text, return the edited text. "
            "If the user wants data analysis or a visualization, describe your analysis "
            "and optionally provide code or steps for creating the requested plot.\n\n"
            f"Instruction: {instruction}\n\n"
            f"Content:\n{initial_content}\n\n"
            "Return the final text or analysis below. If you generate code, enclose it in triple backticks."
        )
        try:
            edited_text = unified_llm.call(prompt)
            result["edited"] = edited_text.strip()
            edit_win.destroy()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to edit/analyze content: {e}")

    submit_btn = ttk.Button(edit_win, text="Submit Instruction", command=submit_edit)
    submit_btn.pack(padx=10, pady=10)

    edit_win.wait_window()
    return result["edited"]

def process_visualization_response(response: str) -> str:
    """
    Processes the LLM response for visualization code.
    If the response contains a code block (delimited by triple backticks),
    it extracts that code, passes it to the visualization runner (verify_and_run_code),
    and replaces the code block with the execution output.
    If a "CHART:" marker is found, it generates a dummy chart.
    Otherwise, returns the original response.
    """
    # Look for code blocks delimited by triple backticks
    code_blocks = re.findall(r"```(.*?)```", response, re.DOTALL)
    if code_blocks:
        code_to_run = code_blocks[0]
        verification_output = verify_and_run_code(code_to_run)
        return response.replace("```" + code_to_run + "```", verification_output)
    if "CHART:" in response:
        # Dummy chart logic; customize as needed.
        chart_type = "BAR"
        x = [1, 2, 3, 4, 5]
        y = [10, 20, 15, 30, 25]
        fig, ax = plt.subplots()
        if chart_type.upper() == "BAR":
            ax.bar(x, y, color='skyblue')
        elif chart_type.upper() == "LINE":
            ax.plot(x, y, marker='o')
        else:
            ax.plot(x, y, marker='o')
        ax.set_title("Generated Chart")
        ax.set_xlabel("X-axis")
        ax.set_ylabel("Y-axis")
        chart_win = tk.Toplevel()
        chart_win.title("Chart")
        canvas = FigureCanvasTkAgg(fig, master=chart_win)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)
        return response.replace("CHART:", "[Chart generated above]")
    return response

def recheck_and_correct_code(code: str) -> str:
    """
    Uses a code analyzer LLM agent to review and correct the provided Python code.
    If the code has syntax or logical errors, it returns a corrected version of the code.
    Otherwise, it returns the original code.
    """
    prompt = (
        "You are a code analyzer. Please review the following Python code for any syntax or logical errors. "
        "If errors are found, provide a corrected version of the code. If the code is correct, return it unchanged.\n\n"
        f"Code:\n{code}\n\n"
        "Corrected Code:"
    )
    corrected_code = unified_llm.call(prompt)
    return corrected_code

def verify_and_run_code(code: str) -> str:
    """
    Verifies and executes visualization code. First, it rechecks and corrects the code using recheck_and_correct_code.
    Then it executes the corrected code and, if a matplotlib figure is generated, displays it in a new window.
    """
    corrected_code = recheck_and_correct_code(code)
    namespace = {}
    try:
        exec(corrected_code, namespace)
        fig = plt.gcf()
        chart_win = tk.Toplevel()
        chart_win.title("Generated Visualization")
        canvas = FigureCanvasTkAgg(fig, master=chart_win)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)
        return "[Visualization generated]"
    except Exception as e:
        logging.error("Error in verify_and_run_code: %s", e)
        return f"[Error executing visualization code: {e}]"

def analyze_uploaded_file(file_path: str) -> str:
    """
    Extracts content from the uploaded file and sends it to the unified LLM
    for analysis or visualization suggestions.
    """
    tool = None
    try:
        tool = GmailDataFetcherTool()
    except Exception as e:
        logging.error("Could not instantiate GmailDataFetcherTool: %s", e)

    if not tool:
        return "Error: GmailDataFetcherTool could not be initialized."

    content = tool.extract_text_from_file(file_path, os.path.basename(file_path))
    prompt = (
        "You are a data analysis assistant. The user has provided file content below. "
        "Please analyze it and suggest any insights, transformations, or potential plots. "
        "If you provide code for a plot, enclose it in triple backticks.\n\n"
        f"File content:\n{content}\n\n"
        "Return your analysis or edited text below."
    )
    try:
        analysis = unified_llm.call(prompt)
        return analysis.strip()
    except Exception as e:
        return f"Error during file analysis: {e}"

class GmailDataFetcherTool(BaseTool):
    name: str = "gmail_data_fetcher"
    description: str = (
        "Fetches Gmail data and attachments, processes emails using an LLM, and allows viewing a summary with "
        "buttons to view full email content and open attachments in a custom editor."
    )
    config: Dict[str, Any] = Field(default_factory=dict)
    conversation_history: List[str] = Field(default_factory=list)
    cache: ClassVar[Dict[str, Any]] = Field(default_factory=dict)

    def __init__(self, **data: Any):
        super().__init__(**data)
        object.__setattr__(self, "config", self.load_external_config("search_config.json"))

    def load_external_config(self, filename: str) -> Dict[str, Any]:
        if os.path.exists(filename):
            with open(filename, "r") as file:
                return json.load(file)
        return {}

    def parse_chunking_method(self, query: str) -> Optional[str]:
        q = query.lower()
        if "chunk by word" in q:
            return "word"
        elif "chunk by sentence" in q:
            return "sentence"
        elif "chunk by paragraph" in q:
            return "paragraph"
        elif "chunk by pages" in q:
            return "pages"
        elif "chunk by batches" in q:
            return "batches"
        elif "chunk by headings_subheadings" in q:
            return "headings_subheadings"
        elif "chunk by headings" in q:
            return "headings"
        else:
            return None

    def advanced_query_parsing(self, query: str) -> str:
        tokens = re.split(r'\W+', query)
        stopwords = self.config.get("stopwords", [])
        filtered_tokens = [token for token in tokens if token and token.lower() not in stopwords]
        synonyms = self.config.get("synonyms", {})
        normalized_tokens = [synonyms.get(token.lower(), token) for token in filtered_tokens]
        parsed = " ".join(normalized_tokens)
        logging.info("Advanced parsed query: %s", parsed)
        return parsed

    def extract_email_text(self, msg: Dict[str, Any], chunk_method: Optional[str] = None,
                             for_summary: bool = True) -> str:
        """Extracts only the plain text body of the email."""
        try:
            email_text = ""
            payload = msg.get('payload', {})
            if 'parts' in payload:
                for part in payload['parts']:
                    if part.get('mimeType', '') == 'text/plain':
                        data = part['body'].get('data')
                        if data:
                            email_text += base64.urlsafe_b64decode(data.encode('UTF-8')).decode('UTF-8')
            else:
                if payload.get('mimeType') == 'text/plain':
                    data = payload['body'].get('data')
                    if data:
                        email_text = base64.urlsafe_b64decode(data.encode('UTF-8')).decode('UTF-8')
            return email_text
        except Exception as e:
            logging.error("Error extracting email text: %s", e)
            return f"Error extracting email text: {e}"

    def get_attachment_info(self, msg: Dict[str, Any]) -> List[Tuple[str, str, Dict[str, Any]]]:
        """Returns a list of attachment info tuples: (filename, message id, part)"""
        attachments = []
        payload = msg.get('payload', {})
        if 'parts' in payload:
            for part in payload['parts']:
                if part.get('filename'):
                    attachments.append((part.get('filename'), msg['id'], part))
        return attachments

    def download_attachment(self, message_id: str, part: Dict[str, Any], download_path: str):
        try:
            service = build('gmail', 'v1', credentials=self.get_gmail_credentials())
            attachment_id = part['body']['attachmentId']
            response = service.users().messages().attachments().get(
                userId='me',
                messageId=message_id,
                id=attachment_id
            ).execute()
            file_data = base64.urlsafe_b64decode(response['data'].encode('UTF-8'))
            file_path = os.path.join(download_path, part['filename'])
            with open(file_path, 'wb') as file:
                file.write(file_data)
            logging.info("Attachment downloaded to %s", file_path)
        except Exception as e:
            logging.error("Error downloading attachment: %s", e)
            return f"Error downloading attachment: {e}"

    def get_gmail_credentials(self):
        creds = None
        token_path = 'token.pickle'
        if os.path.exists(token_path):
            with open(token_path, 'rb') as token:
                creds = pickle.load(token)
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
                creds = flow.run_local_server(port=0)
            with open(token_path, 'wb') as token:
                pickle.dump(creds, token)
        return creds

    def analyze_attachment(self, message_id: str, part: Dict[str, Any]) -> str:
        """Extracts full content from the attachment for analysis."""
        try:
            temp_dir = tempfile.mkdtemp()
            self.download_attachment(message_id, part, temp_dir)
            file_path = os.path.join(temp_dir, part.get('filename', 'unknown'))
            content = self.extract_text_from_file(file_path, os.path.basename(file_path))
            return content
        except Exception as e:
            logging.error("Error analyzing attachment: %s", e)
            return f"Error analyzing attachment: {e}"

    def extract_text_from_file(self, file_path: str, file_name: str) -> str:
        text = ""
        fname = file_name.lower()
        try:
            if fname.endswith('.pdf'):
                with open(file_path, 'rb') as file_content:
                    reader = PdfReader(file_content)
                    text = "".join(p.extract_text() for p in reader.pages if p.extract_text())
            elif fname.endswith(('.xlsx', '.xls')):
                workbook = openpyxl.load_workbook(file_path)
                sheet = workbook.active
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    rows.append("\t".join(str(cell) if cell is not None else "" for cell in row))
                text = "\n".join(rows)
            elif fname.endswith('.docx'):
                document = docx.Document(file_path)
                text = "\n".join(p.text for p in document.paragraphs)
            elif fname.endswith('.csv'):
                with open(file_path, 'r', encoding='utf-8') as file_content:
                    reader = csv.reader(file_content)
                    rows = []
                    for row in reader:
                        rows.append("\t".join(row))
                    text = "\n".join(rows)
            elif fname.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                if pytesseract:
                    try:
                        image = Image.open(file_path)
                        text = pytesseract.image_to_string(image)
                    except Exception as e:
                        text = f"Error during OCR: {e}"
                else:
                    text = "OCR support not installed for image extraction."
            else:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            return text
        except Exception as e:
            logging.error("Error extracting text from %s: %s", file_name, e)
            return f"Error extracting text from {file_name}: {e}"

    def cleanup_temp_dir(self, temp_dir: str):
        import shutil
        shutil.rmtree(temp_dir)

    def _run(self, query: str) -> List[Dict[str, Any]]:
        """
        Standard email query entry point.
        For each email, extracts the email body (ignoring attachment content), computes a summary,
        and records attachment filenames.
        Returns a list of email info dictionaries:
            { "summary": ..., "full": ..., "attachments": [...] }
        """
        try:
            logging.info(f"Running GmailDataFetcherTool with query: {query}")
            intent = interpret_user_query(query)
            logging.info(f"Interpreted intent: {intent}")
            gmail_query = intent.get("gmail_query", query)
            chunk_method = self.parse_chunking_method(query)
            logging.info(f"Chunk method: {chunk_method}")
            messages = fetch_gmail_data(query=gmail_query, batch_size=10)
            logging.info(f"Fetched {len(messages)} messages.")
            email_info_list = []
            for msg in messages:
                full_text = self.extract_email_text(msg, chunk_method, for_summary=True)
                summary_prompt = f"Summarize the following email content in one or two sentences:\n{full_text}"
                summary = llm_predict_with_backoff(summary_prompt)
                attachments = self.get_attachment_info(msg)
                email_info_list.append({
                    "summary": summary,
                    "full": full_text,
                    "attachments": attachments
                })
            return email_info_list
        except Exception as e:
            logging.error(f"Tool execution failed: {e}", exc_info=True)
            return [{"summary": f"Error processing Gmail data: {e}", "full": "", "attachments": []}]

    def open_attachment(self, attachment_info: Tuple[str, str, Dict[str, Any]]) -> None:
        """Download the attachment and open it in our custom editor."""
        filename, message_id, part = attachment_info
        temp_dir = tempfile.mkdtemp()
        self.download_attachment(message_id, part, temp_dir)
        file_path = os.path.join(temp_dir, filename)
        open_custom_editor(file_path)

def custom_chunk_text(text: str, method: str) -> List[str]:
    """Splits text into chunks based on the specified method."""
    if method == "word":
        words = text.split()
        return [" ".join(words[i:i+50]) for i in range(0, len(words), 50)]
    elif method == "sentence":
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return sentences
    elif method == "paragraph":
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        return paragraphs
    elif method == "pages":
        pages = re.split(r'\n{2,}', text)
        return pages
    elif method == "batches":
        pages = re.split(r'\n{2,}', text)
        batch_size = 3
        batches = []
        for i in range(0, len(pages), batch_size):
            batches.append("\n\n".join(pages[i:i+batch_size]))
        return batches
    elif method == "headings":
        lines = text.splitlines()
        chunks = []
        current_chunk = []
        for line in lines:
            if line.isupper() and len(line.strip()) > 3:
                if current_chunk:
                    chunks.append("\n".join(current_chunk))
                    current_chunk = []
                current_chunk.append(line)
            else:
                current_chunk.append(line)
        if current_chunk:
            chunks.append("\n".join(current_chunk))
        return chunks
    elif method == "headings_subheadings":
        lines = text.splitlines()
        chunks = []
        current_chunk = []
        for line in lines:
            if (line.isupper() and len(line.strip()) > 3) or (line.istitle() and len(line.split()) < 5):
                if current_chunk:
                    chunks.append("\n".join(current_chunk))
                    current_chunk = []
                current_chunk.append(line)
            else:
                current_chunk.append(line)
        if current_chunk:
            chunks.append("\n".join(current_chunk))
        return chunks
    else:
        return [text]

def visualize_attachment_file(file_path: str) -> None:
    """
    Reads the file content from the given attachment file, sends a prompt to the LLM
    to generate Python visualization code (using matplotlib, seaborn, plotly, or graphviz),
    rechecks and corrects the code via the code analyzer LLM (using recheck_and_correct_code),
    and then executes the corrected code to generate and display visualizations.
    """
    content = extract_text_from_file(file_path)
    prompt = (
        "You are an expert data analyst and visualization assistant. "
        "Given the following file content, generate Python code (using matplotlib, seaborn, plotly, or graphviz) "
        "that produces 5 key charts providing the most important insights from the data. "
        "Return your code enclosed in triple backticks.\n\n"
        f"File content:\n{content}\n\n"
        "Visualization Code:"
    )
    response = unified_llm.call(prompt)
    # Process the response to extract and run the visualization code.
    final_response = process_visualization_response(response)
    # Optionally, display the final response for debugging or confirmation.
    messagebox.showinfo("Visualization Result", final_response)

def open_custom_editor_attachment(attachment_info: Tuple[str, str, Dict[str, Any]]) -> None:
    """
    Downloads the attachment and then calls visualize_attachment_file
    to generate and display visualizations based on the attachment's content.
    """
    filename, message_id, part = attachment_info
    temp_dir = tempfile.mkdtemp()
    GmailDataFetcherTool().download_attachment(message_id, part, temp_dir)
    file_path = os.path.join(temp_dir, filename)
    # Instead of opening the text editor, send the file for visualization.
    visualize_attachment_file(file_path)

def open_custom_editor(file_path: str) -> None:
    """
    Opens a custom editor window to display and allow editing of the content of the given file.
    For structured data (Excel/CSV): displays content in a table.
    For PDFs/DOCX/TXT: displays extracted text.
    For images: displays the image with pan/zoom and provides commands for OCR/analysis.
    A natural language command box allows editing via an AI agent.
    """
    editor = tk.Toplevel()
    editor.title(f"Custom Editor - {os.path.basename(file_path)}")
    editor.geometry("900x700")
    ext = os.path.splitext(file_path)[1].lower()

    if ext in [".xlsx", ".xls", ".csv"]:
        # Structured data: display in a table using a Treeview.
        frame = ttk.Frame(editor)
        frame.pack(fill="both", expand=True, padx=10, pady=10)
        tree = ttk.Treeview(frame)
        tree.pack(side="left", fill="both", expand=True)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=tree.yview)
        scrollbar.pack(side="right", fill="y")
        tree.configure(yscrollcommand=scrollbar.set)
        data = []
        headers = []
        if ext in [".xlsx", ".xls"]:
            try:
                workbook = openpyxl.load_workbook(file_path)
                sheet = workbook.active
                for i, row in enumerate(sheet.iter_rows(values_only=True)):
                    if i == 0:
                        headers = [str(cell) if cell is not None else "" for cell in row]
                    else:
                        data.append([str(cell) if cell is not None else "" for cell in row])
            except Exception as e:
                headers = ["Error"]
                data = [[f"Error loading Excel: {e}"]]
        elif ext == ".csv":
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    reader = csv.reader(f)
                    for i, row in enumerate(reader):
                        if i == 0:
                            headers = row
                        else:
                            data.append(row)
            except Exception as e:
                headers = ["Error"]
                data = [[f"Error loading CSV: {e}"]]
        tree["columns"] = headers
        tree["show"] = "headings"
        for header in headers:
            tree.heading(header, text=header)
            tree.column(header, width=100)
        for row in data:
            tree.insert("", "end", values=row)
        # Chunking drop-down for table data (convert table to text for chunking)
        chunk_methods = ["word", "sentence", "paragraph", "pages", "batches", "headings", "headings_subheadings"]
        chunk_var = tk.StringVar(value="paragraph")
        ttk.Label(editor, text="Select Chunking Method:").pack(anchor="w", padx=10, pady=(10, 0))
        chunk_dropdown = ttk.Combobox(editor, textvariable=chunk_var, values=chunk_methods, state="readonly")
        chunk_dropdown.pack(anchor="w", padx=10, pady=5)

        def chunk_content_table():
            table_text = "\n".join(["\t".join(headers)] + ["\t".join(row) for row in data])
            method = chunk_var.get()
            chunks = custom_chunk_text(table_text, method)
            new_text = "\n--- Chunk ---\n".join(chunks)
            # Show chunked text in a new ScrolledText widget.
            for widget in editor.winfo_children():
                widget.destroy()
            new_text_widget = ScrolledText(editor, wrap=tk.WORD)
            new_text_widget.insert("1.0", new_text)
            new_text_widget.pack(fill="both", expand=True, padx=10, pady=10)

        chunk_btn = ttk.Button(editor, text="Chunk Content", command=chunk_content_table)
        chunk_btn.pack(anchor="w", padx=10, pady=5)
    elif ext in [".pdf", ".docx", ".txt"]:
        # For PDFs, DOCX, and TXT: display extracted text.
        content = ""
        try:
            tool = GmailDataFetcherTool()
            content = tool.extract_text_from_file(file_path, os.path.basename(file_path))
        except Exception as e:
            content = f"Error loading content: {e}"
        content_widget = ScrolledText(editor, wrap=tk.WORD, height=25)
        content_widget.insert("1.0", content)
        content_widget.pack(fill="both", expand=True, padx=10, pady=10)
        # Drop-down for custom chunking.
        chunk_methods = ["word", "sentence", "paragraph", "pages", "batches", "headings", "headings_subheadings"]
        chunk_var = tk.StringVar(value="paragraph")
        ttk.Label(editor, text="Select Chunking Method:").pack(anchor="w", padx=10, pady=(10, 0))
        chunk_dropdown = ttk.Combobox(editor, textvariable=chunk_var, values=chunk_methods, state="readonly")
        chunk_dropdown.pack(anchor="w", padx=10, pady=5)

        def chunk_content():
            current_content = content_widget.get("1.0", tk.END)
            method = chunk_var.get()
            chunks = custom_chunk_text(current_content, method)
            new_text = "\n--- Chunk ---\n".join(chunks)
            content_widget.delete("1.0", tk.END)
            content_widget.insert("1.0", new_text)

        chunk_btn = ttk.Button(editor, text="Chunk Content", command=chunk_content)
        chunk_btn.pack(anchor="w", padx=10, pady=5)
    elif ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp"]:
        # Image: display in canvas with pan/zoom.
        canvas = tk.Canvas(editor, bg="gray")
        canvas.pack(fill="both", expand=True)
        try:
            image = Image.open(file_path)
            photo = ImageTk.PhotoImage(image)
            canvas.create_image(0, 0, anchor="nw", image=photo)
            canvas.image = photo
            hbar = tk.Scrollbar(editor, orient=tk.HORIZONTAL, command=canvas.xview)
            hbar.pack(side=tk.BOTTOM, fill=tk.X)
            vbar = tk.Scrollbar(editor, orient=tk.VERTICAL, command=canvas.yview)
            vbar.pack(side=tk.RIGHT, fill=tk.Y)
            canvas.config(xscrollcommand=hbar.set, yscrollcommand=vbar.set, scrollregion=canvas.bbox(tk.ALL))

            def zoom(event):
                factor = 1.001 ** event.delta
                canvas.scale(tk.ALL, event.x, event.y, factor, factor)
                canvas.configure(scrollregion=canvas.bbox(tk.ALL))

            canvas.bind("<MouseWheel>", zoom)
        except Exception as e:
            tk.Label(editor, text=f"Error loading image: {e}").pack()
        command_label = ttk.Label(editor, text="Enter command for image analysis (e.g., 'Extract text'):")
        command_label.pack(anchor="w", padx=10, pady=(10, 0))
        command_entry = ttk.Entry(editor, width=80)
        command_entry.pack(anchor="w", padx=10, pady=5)

        def process_command_image():
            command = command_entry.get().strip()
            if not command:
                messagebox.showerror("Error", "Command cannot be empty.")
                return
            if "extract text" in command.lower():
                try:
                    ocr_text = pytesseract.image_to_string(image)
                except Exception as e:
                    ocr_text = f"Error during OCR: {e}"
                messagebox.showinfo("OCR Result", ocr_text)
            else:
                prompt = (
                    "You are an image analysis assistant. The user has provided an image and a command.\n"
                    f"Command: {command}\n"
                    "Return your analysis."
                )
                analysis = unified_llm.call(prompt)
                messagebox.showinfo("Image Analysis", analysis)

        process_btn = ttk.Button(editor, text="Process Command", command=process_command_image)
        process_btn.pack(anchor="w", padx=10, pady=5)
    else:
        messagebox.showerror("Error", "Unknown file format.")
        return

###############################################################################
# GmailDataFetcherUI
###############################################################################
class GmailDataFetcherUI:
    def __init__(self, root: tk.Tk):
        self.root = root
        self.root.title("Gmail Data Fetcher")
        self.fetcher_tool = GmailDataFetcherTool()

        # Input frame for Gmail query
        self.input_frame = ttk.Frame(self.root, padding="10")
        self.input_frame.pack(fill="x")

        self.query_label = ttk.Label(self.input_frame, text="Enter Gmail Query:", font=("Arial", 12))
        self.query_label.pack(side="left")

        self.query_entry = ttk.Entry(self.input_frame, width=50, font=("Arial", 12))
        self.query_entry.pack(side="left", padx=5)

        self.fetch_button = ttk.Button(self.input_frame, text="Fetch and Process", command=self.fetch_and_process)
        self.fetch_button.pack(side="left", padx=5)

        self.advanced_analysis_var = tk.BooleanVar(value=False)
        self.advanced_check = ttk.Checkbutton(
            self.input_frame, text="Enable Advanced Analysis",
            variable=self.advanced_analysis_var
        )
        self.advanced_check.pack(side="left", padx=5)

        # File Upload frame for analysis
        self.upload_frame = ttk.LabelFrame(self.root, text="Upload Files", padding="10")
        self.upload_frame.pack(fill="x", padx=10, pady=5)
        self.upload_label = ttk.Label(self.upload_frame, text="Drag & Drop or Upload File for Analysis:", font=("Arial", 12))
        self.upload_label.pack(side="left")
        self.upload_btn = ttk.Button(self.upload_frame, text="Upload File", command=self.upload_and_analyze_file)
        self.upload_btn.pack(side="left", padx=5)

        # Emails panel
        self.emails_frame = ttk.Frame(self.root, padding="10")
        self.emails_frame.pack(fill="both", expand=True, padx=10, pady=10)

        # Status bar
        self.status_bar = ttk.Label(self.root, text="Ready", relief=tk.SUNKEN, anchor=tk.W, padding="5")
        self.status_bar.pack(side="bottom", fill="x")

    def fetch_and_process(self):
        query = self.query_entry.get()
        if not query:
            messagebox.showerror("Error", "Query cannot be empty.")
            return

        self.status_bar.config(text="Fetching and processing...")
        self.root.update_idletasks()

        if query.lower().startswith("analyze "):
            file_name = query[len("analyze "):].strip()
            if self.advanced_analysis_var.get():
                result = self.fetcher_tool.advanced_analysis(file_name)
                self.open_analysis_window(result)
            else:
                messagebox.showinfo("Advanced Analysis Disabled", "Please enable Advanced Analysis to proceed.")
            self.status_bar.config(text="Done.")
            return

        try:
            intent = interpret_user_query(query)
            logging.info(f"Interpreted intent: {intent}")
            gmail_query = intent.get("gmail_query", query)
            email_info_list = self.fetcher_tool._run(gmail_query)
            for widget in self.emails_frame.winfo_children():
                widget.destroy()
            for email_info in email_info_list:
                frame = ttk.Frame(self.emails_frame, relief=tk.RIDGE, borderwidth=1)
                frame.pack(fill="x", pady=5)
                summary_label = ttk.Label(frame, text="Summary: " + email_info["summary"], wraplength=600, justify=tk.LEFT)
                summary_label.pack(side="top", anchor="w", padx=5, pady=2)
                btn_full = ttk.Button(frame, text="View Full Email", command=lambda text=email_info["full"]: self.open_full_email(text))
                btn_full.pack(side="top", anchor="w", padx=5, pady=2)
                if email_info["attachments"]:
                    attach_label = ttk.Label(frame, text="Attachments:")
                    attach_label.pack(side="top", anchor="w", padx=5)
                    for attachment_info in email_info["attachments"]:
                        filename, _, _ = attachment_info
                        btn_attach = ttk.Button(frame, text=f"Open {filename}", command=lambda ai=attachment_info: self.open_attachment(ai))
                        btn_attach.pack(side="top", anchor="w", padx=15, pady=2)
            self.status_bar.config(text="Done.")
        except Exception as e:
            messagebox.showerror("Error", str(e))
            self.status_bar.config(text=f"Error: {e}")
        finally:
            self.root.update_idletasks()

    def open_full_email(self, full_text: str):
        top = tk.Toplevel(self.root)
        top.title("Full Email Content")
        text_widget = ScrolledText(top, wrap=tk.WORD, height=25, width=80)
        text_widget.pack(fill="both", expand=True, padx=10, pady=10)
        text_widget.insert("1.0", full_text)

    def open_attachment(self, attachment_info: Tuple[str, str, Dict[str, Any]]):
        open_custom_editor_attachment(attachment_info)

    def open_analysis_window(self, analysis_result: str):
        top = tk.Toplevel(self.root)
        top.title("Advanced Analysis")
        analysis_text = ScrolledText(top, wrap=tk.WORD, height=20, width=80)
        analysis_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        analysis_text.insert("1.0", analysis_result)

    def upload_and_analyze_file(self):
        file_path = filedialog.askopenfilename(
            title="Select a file for analysis",
            filetypes=[
                ("All Files", "*.*"),
                ("PDF", "*.pdf"),
                ("Image", "*.png;*.jpg;*.jpeg;*.gif;*.bmp"),
                ("CSV", "*.csv"),
                ("Excel", "*.xlsx;*.xls"),
                ("Word/Text", "*.docx;*.txt")
            ]
        )
        if file_path:
            analysis = analyze_uploaded_file(file_path)
            top = tk.Toplevel(self.root)
            top.title("File Analysis Result")
            result_text = ScrolledText(top, wrap=tk.WORD, height=25, width=80)
            result_text.pack(fill="both", expand=True, padx=10, pady=10)
            result_text.insert("1.0", analysis)

###############################################################################
# MultiDocRAGApp and related classes
###############################################################################
class DocumentStore:
    def __init__(self, use_embeddings: bool = True, chunk_size: int = 500):
        self.use_embeddings = use_embeddings and (EMBEDDING_MODEL is not None)
        self.chunk_size = chunk_size
        self.docs: List[Tuple[str, str]] = []  # (short_name, chunk)
        self.embeddings = []  # Embeddings for each chunk
        self.short_to_original: Dict[str, str] = {}
        self.original_to_short: Dict[str, str] = {}
        self.file_counter = 1

    def _generate_short_name(self, original_name: str) -> str:
        short_name = f"File_{self.file_counter}"
        self.file_counter += 1
        self.short_to_original[short_name] = original_name
        self.original_to_short[original_name] = short_name
        return short_name

    def add_document_text(self, text: str, original_name: str):
        short_name = self._generate_short_name(original_name)
        chunks = self._chunk_text(text, self.chunk_size)
        for chunk in chunks:
            self.docs.append((short_name, chunk))
            if self.use_embeddings:
                emb = EMBEDDING_MODEL.encode(chunk, convert_to_tensor=True)
                self.embeddings.append(emb)

    def _chunk_text(self, text: str, chunk_size: int) -> List[str]:
        chunks = []
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunks.append(text[start:end])
            start = end
        return chunks

    def retrieve_relevant_chunks(self, query: str, top_k: int = 3) -> List[str]:
        # Fuzzy matching to find a referenced file
        if fuzz is not None:
            best_score = 0
            best_original = None
            for orig in self.original_to_short:
                score = fuzz.partial_ratio(orig.lower(), query.lower())
                if score > best_score:
                    best_score = score
                    best_original = orig
            if best_score > 70:
                short_name = self.original_to_short[best_original]
                return self._retrieve_from_short_name(short_name, query, top_k)
        else:
            for orig in self.original_to_short:
                if orig.lower() in query.lower():
                    short_name = self.original_to_short[orig]
                    return self._retrieve_from_short_name(short_name, query, top_k)
        # Fallback: retrieve from all docs
        if self.use_embeddings:
            query_emb = EMBEDDING_MODEL.encode(query, convert_to_tensor=True)
            scores = util.pytorch_cos_sim(query_emb, self.embeddings)[0]
            top_results = scores.topk(k=min(top_k, len(scores))).indices.tolist()
            return [f"From file '{self.docs[idx][0]}':\n{self.docs[idx][1]}" for idx in top_results]
        else:
            tokens_query = set(query.lower().split())
            def score(item: Tuple[str, str]) -> int:
                return len(tokens_query.intersection(set(item[1].lower().split())))
            ranked = sorted(self.docs, key=score, reverse=True)
            return [f"From file '{doc[0]}':\n{doc[1]}" for doc in ranked[:top_k]]

    def _retrieve_from_short_name(self, short_name: str, query: str, top_k: int) -> List[str]:
        relevant_docs = [(sn, chunk) for (sn, chunk) in self.docs if sn == short_name]
        if not relevant_docs:
            return [f"No chunks found for file {short_name}"]
        if self.use_embeddings:
            relevant_indices = [i for i, (sn, _) in enumerate(self.docs) if sn == short_name]
            relevant_chunks = [self.docs[i][1] for i in relevant_indices]
            relevant_embs = [self.embeddings[i] for i in relevant_indices]
            query_emb = EMBEDDING_MODEL.encode(query, convert_to_tensor=True)
            scores = util.pytorch_cos_sim(query_emb, relevant_embs)[0]
            top_results = scores.topk(k=min(top_k, len(scores))).indices.tolist()
            return [f"From file '{short_name}':\n{relevant_chunks[idx]}" for idx in top_results]
        else:
            tokens_query = set(query.lower().split())
            def score(item: Tuple[str, str]) -> int:
                return len(tokens_query.intersection(set(item[1].lower().split())))
            ranked = sorted(relevant_docs, key=score, reverse=True)
            return [f"From file '{short_name}':\n{doc[1]}" for doc in ranked[:top_k]]

class RAGChatBot:
    def __init__(self, doc_store: DocumentStore):
        self.doc_store = doc_store
        self.conversation_history = ""

    def add_document(self, text: str, original_name: str):
        self.doc_store.add_document_text(text, original_name)

    def chat(self, user_query: str, top_k: int = 3) -> str:
        relevant_chunks = self.doc_store.retrieve_relevant_chunks(user_query, top_k=top_k)
        context_block = "\n\n".join([f"Relevant chunk:\n{chunk}" for chunk in relevant_chunks])
        prompt = (
            "You are a helpful assistant with access to local documents via the doc store. "
            "If the user references a file by name, use that file's data.\n\n"
            f"{context_block}\n\n"
            f"User: {user_query}\nAssistant:"
        )
        response = unified_llm.call(prompt)
        self.conversation_history += f"\nUser: {user_query}\nAssistant: {response}"
        return process_visualization_response(response)

class MultiDocRAGApp:
    def __init__(self, parent: tk.Widget):
        self.parent = parent
        self.doc_store = DocumentStore(use_embeddings=True, chunk_size=500)
        self.chatbot = RAGChatBot(self.doc_store)

        self.upload_frame = ttk.LabelFrame(self.parent, text="Upload Files")
        self.upload_frame.pack(fill="x", padx=10, pady=5)
        self.upload_button = ttk.Button(self.upload_frame, text="Add Files", command=self.add_files)
        self.upload_button.pack(side="left", padx=5, pady=5)
        self.file_listbox = tk.Listbox(self.upload_frame, width=60)
        self.file_listbox.pack(side="left", fill="both", expand=True, padx=5, pady=5)

        self.chat_frame = ttk.LabelFrame(self.parent, text="Chat with Your Documents")
        self.chat_frame.pack(fill="both", expand=True, padx=10, pady=5)
        self.chat_display = ScrolledText(self.chat_frame, wrap=tk.WORD, height=15)
        self.chat_display.pack(fill="both", expand=True, padx=5, pady=5)

        self.user_input_frame = ttk.Frame(self.chat_frame)
        self.user_input_frame.pack(fill="x", padx=5, pady=5)
        self.user_entry = tk.Entry(self.user_input_frame, width=60)
        self.user_entry.pack(side="left", fill="x", expand=True, padx=5)
        self.send_button = ttk.Button(self.user_input_frame, text="Send", command=self.send_query)
        self.send_button.pack(side="left", padx=5)
        self.clear_history_btn = ttk.Button(self.user_input_frame, text="Clear Chat History", command=self.clear_history)
        self.clear_history_btn.pack(side="left", padx=5)
        self.show_history_btn = ttk.Button(self.user_input_frame, text="Show Full History", command=self.show_history)
        self.show_history_btn.pack(side="left", padx=5)
        self.status_bar = ttk.Label(self.parent, text="Ready", relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(side="bottom", fill="x")

    def add_files(self):
        file_paths = filedialog.askopenfilenames(
            title="Select files",
            filetypes=[
                ("All Files", "*.*"),
                ("PDF", "*.pdf"),
                ("Image", "*.png;*.jpg;*.jpeg;*.gif;*.bmp"),
                ("CSV", "*.csv"),
                ("Excel", "*.xlsx;*.xls"),
                ("Word/Text", "*.docx;*.txt")
            ]
        )
        for path in file_paths:
            self.file_listbox.insert(tk.END, path)
            self.status_bar.config(text=f"Processing {os.path.basename(path)}...")
            self.parent.update_idletasks()
            text = extract_text_from_file(path)
            self.chatbot.add_document(text, os.path.basename(path))
        self.status_bar.config(text="Files added successfully.")

    def send_query(self):
        query = self.user_entry.get().strip()
        if not query:
            return
        self.user_entry.delete(0, tk.END)
        self.chat_display.insert(tk.END, f"\nUser: {query}\n")
        self.chat_display.see(tk.END)
        threading.Thread(target=self._process_query, args=(query,)).start()

    def _process_query(self, query: str):
        self.status_bar.config(text="Thinking...")
        self.parent.update_idletasks()
        try:
            response = self.chatbot.chat(query)
            self.chat_display.insert(tk.END, f"Assistant: {response}\n")
            self.chat_display.see(tk.END)
        except Exception as e:
            self.chat_display.insert(tk.END, f"Error: {e}\n")
        finally:
            self.status_bar.config(text="Ready")

    def clear_history(self):
        self.chatbot.conversation_history = ""
        self.chat_display.delete("1.0", tk.END)
        messagebox.showinfo("Chat History Cleared", "The conversation history has been cleared.")

    def show_history(self):
        history_win = tk.Toplevel()
        history_win.title("Full Conversation History")
        history_text = ScrolledText(history_win, wrap=tk.WORD, height=20, width=80)
        history_text.pack(fill="both", expand=True, padx=10, pady=10)
        history_text.insert("1.0", self.chatbot.conversation_history)

class AICodeEditorUI:
    def __init__(self, parent: tk.Widget):
        self.parent = parent
        self.code_frame = ttk.LabelFrame(self.parent, text="AI Code Editor")
        self.code_frame.pack(fill="both", expand=True, padx=10, pady=5)

        self.code_entry = ScrolledText(self.code_frame, wrap=tk.WORD, height=20)
        self.code_entry.pack(fill="both", expand=True, padx=5, pady=5)

        self.run_button = ttk.Button(self.code_frame, text="Run Code", command=self.run_code)
        self.run_button.pack(side="left", padx=5, pady=5)

        self.output_frame = ttk.LabelFrame(self.parent, text="Output")
        self.output_frame.pack(fill="both", expand=True, padx=10, pady=5)

        self.output_display = ScrolledText(self.output_frame, wrap=tk.WORD, height=10)
        self.output_display.pack(fill="both", expand=True, padx=5, pady=5)

    def run_code(self):
        code = self.code_entry.get("1.0", tk.END).strip()
        if not code:
            messagebox.showerror("Error", "Code cannot be empty.")
            return
        try:
            output = verify_and_run_code(code)
            self.output_display.delete("1.0", tk.END)
            self.output_display.insert(tk.END, output)
        except Exception as e:
            messagebox.showerror("Error", f"Error executing code: {e}")

class MainApp(ttk.Frame):
    def __init__(self, root: tk.Tk):
        super().__init__(root)
        root.title("Unified Application: Gmail & Document Analysis + AI Code Editor")
        root.geometry("1000x800")
        root.configure(bg="#f0f0f0")  # Light background color

        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill="both", expand=True)

        # Gmail Data Tab
        self.gmail_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.gmail_frame, text="Gmail Data")
        self.gmail_ui = GmailDataFetcherUI(root)

        # Document Chat Tab
        self.doc_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.doc_frame, text="Document Chat")
        self.doc_ui = MultiDocRAGApp(self.doc_frame)

        # AI Code Editor Tab
        self.code_editor_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.code_editor_frame, text="AI Code Editor")
        self.code_editor_ui = AICodeEditorUI(self.code_editor_frame)

def main():
    root = tk.Tk()
    app = MainApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
